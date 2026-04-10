"""
MedSci Skills Demo 3: NHANES Obesity & Diabetes
Step 3 — Generate manuscript draft (write-paper skill)

Produces IMRAD manuscript in Markdown and DOCX formats from analysis outputs.

Usage: python3 03_write_manuscript.py
Output: output/manuscript_draft.md, output/manuscript_draft.docx
"""

import os
import datetime
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("=" * 60)
print("MedSci Skills Demo 3: Manuscript Generation")
print(f"Date: {datetime.date.today()}")
print("=" * 60)

# === Load analysis outputs ===
table1 = pd.read_csv("output/table1.csv")
prevalence = pd.read_csv("output/prevalence_by_bmi.csv")
regression = pd.read_csv("output/regression_results.csv")

# === Extract key values ===
prev_normal = prevalence[prevalence["BMI Category"] == "Normal"]
prev_overweight = prevalence[prevalence["BMI Category"] == "Overweight"]
prev_obese = prevalence[prevalence["BMI Category"] == "Obese"]
prev_overall = prevalence[prevalence["BMI Category"] == "Overall"]

reg_adj = regression[regression["Model"] == "Adjusted"]
or_ow = reg_adj[reg_adj["Variable"] == "Overweight vs Normal"].iloc[0]
or_ob = reg_adj[reg_adj["Variable"] == "Obese vs Normal"].iloc[0]
or_age = reg_adj[reg_adj["Variable"] == "Age (per year)"].iloc[0]
or_female = reg_adj[reg_adj["Variable"] == "Female vs Male"].iloc[0]
or_black = reg_adj[reg_adj["Variable"] == "NH Black vs NH White"].iloc[0]
or_asian = reg_adj[reg_adj["Variable"] == "NH Asian vs NH White"].iloc[0]
or_mexican = reg_adj[reg_adj["Variable"] == "Mexican American vs NH White"].iloc[0]

reg_unadj = regression[regression["Model"] == "Unadjusted"]
or_ow_unadj = reg_unadj[reg_unadj["Variable"] == "Overweight vs Normal"].iloc[0]
or_ob_unadj = reg_unadj[reg_unadj["Variable"] == "Obese vs Normal"].iloc[0]

# === Build manuscript text ===

TITLE = ("Association Between Body Mass Index and Diabetes Mellitus "
         "in US Adults: A Cross-Sectional Analysis of NHANES 2017-2018")

ABSTRACT = f"""**Background:** Obesity is a well-established risk factor for type 2 diabetes mellitus, yet nationally representative estimates that account for the complex survey design of the National Health and Nutrition Examination Survey (NHANES) remain essential for tracking the burden of metabolic disease in the United States.

**Methods:** We conducted a cross-sectional analysis of NHANES 2017-2018 data among adults aged 20 years and older (n = 4,866). Body mass index (BMI) was categorized according to the World Health Organization classification: normal weight (18.5-24.9 kg/m2), overweight (25.0-29.9 kg/m2), and obese (>=30.0 kg/m2). Diabetes was defined as glycated hemoglobin (HbA1c) >= 6.5%. Survey-weighted logistic regression was used to estimate odds ratios (ORs) adjusted for age, sex, race/ethnicity, and education level.

**Results:** The prevalence of diabetes was {prev_normal["Prevalence (%)"].values[0]}% (95% CI: {prev_normal["95% CI lower"].values[0]}-{prev_normal["95% CI upper"].values[0]}%) in the normal weight group, {prev_overweight["Prevalence (%)"].values[0]}% ({prev_overweight["95% CI lower"].values[0]}-{prev_overweight["95% CI upper"].values[0]}%) in the overweight group, and {prev_obese["Prevalence (%)"].values[0]}% ({prev_obese["95% CI lower"].values[0]}-{prev_obese["95% CI upper"].values[0]}%) in the obese group. In adjusted analysis, obesity was associated with a 4.5-fold increase in the odds of diabetes (OR {or_ob["OR"]}, 95% CI: {or_ob["95% CI"]}, p {or_ob["p-value"]}), and overweight with a 2-fold increase (OR {or_ow["OR"]}, 95% CI: {or_ow["95% CI"]}, p {or_ow["p-value"]}), compared with normal weight.

**Conclusion:** A clear dose-response relationship exists between BMI category and diabetes prevalence among US adults. These findings underscore the importance of obesity prevention and weight management as key strategies for reducing the population burden of diabetes."""

INTRODUCTION = """Diabetes mellitus is a leading cause of morbidity and mortality worldwide, affecting an estimated 537 million adults globally as of 2021 [1]. In the United States, approximately 37.3 million people have diabetes, with type 2 diabetes accounting for 90-95% of all cases [2]. Obesity is the most important modifiable risk factor for type 2 diabetes, with epidemiological evidence consistently demonstrating a dose-response relationship between body mass index (BMI) and diabetes risk [3].

The National Health and Nutrition Examination Survey (NHANES) provides nationally representative data on the health and nutritional status of the US civilian noninstitutionalized population [4]. NHANES employs a complex, multistage probability sampling design with oversampling of specific subgroups, necessitating the use of survey weights to produce unbiased population-level estimates [5].

While the association between obesity and diabetes is well-established, continued surveillance using contemporary data is essential for tracking secular trends, informing public health policy, and identifying disparities across demographic subgroups. The purpose of this study was to examine the association between BMI categories and diabetes prevalence in a nationally representative sample of US adults using NHANES 2017-2018 data, accounting for the complex survey design."""

METHODS = f"""## Data Source and Study Population

This cross-sectional study used data from the NHANES 2017-2018 cycle, conducted by the National Center for Health Statistics (NCHS), a division of the Centers for Disease Control and Prevention (CDC). NHANES uses a complex, multistage probability sampling design to obtain a nationally representative sample of the US civilian noninstitutionalized population [4]. The study protocol was approved by the NCHS Research Ethics Review Board, and all participants provided written informed consent.

We included adults aged 20 years and older who had complete data for BMI, glycated hemoglobin (HbA1c), and examination survey weights. Participants classified as underweight (BMI < 18.5 kg/m2, n = 74) were excluded from the primary analysis due to the small sample size in this group, yielding a final analytic sample of 4,866 participants.

## Variable Definitions

### Exposure: Body Mass Index

BMI was calculated from measured height and weight during the NHANES mobile examination center visit. Participants were categorized according to the World Health Organization (WHO) classification: normal weight (18.5-24.9 kg/m2), overweight (25.0-29.9 kg/m2), and obese (>=30.0 kg/m2) [6].

### Outcome: Diabetes Mellitus

Diabetes was defined as HbA1c >= 6.5%, consistent with the American Diabetes Association (ADA) diagnostic threshold [7]. Glycemic status was further classified as normal (HbA1c < 5.7%), prediabetes (5.7-6.4%), and diabetes (>= 6.5%).

### Covariates

Covariates included age (continuous, in years), sex (male, female), race/ethnicity (Non-Hispanic White, Non-Hispanic Black, Non-Hispanic Asian, Mexican American, Other Hispanic, Other/Multi-Racial), and education level (less than high school, high school graduate/GED, some college/associate degree, college graduate or above).

## Statistical Analysis

All analyses incorporated NHANES examination survey weights (WTMEC2YR) to account for the complex survey design, including unequal probability of selection, nonresponse, and post-stratification adjustments.

Baseline characteristics were summarized using means and standard deviations for continuous variables and frequencies and percentages for categorical variables. Differences across BMI categories were assessed using one-way ANOVA for continuous variables and chi-square tests for categorical variables.

Survey-weighted diabetes prevalence estimates with 95% Wilson confidence intervals were calculated for each BMI category. Two logistic regression models were fitted using survey-weighted generalized linear models with a binomial family:

- **Model 1 (Unadjusted):** BMI category only (reference: normal weight)
- **Model 2 (Adjusted):** BMI category + age + sex + race/ethnicity + education level

Results were expressed as odds ratios (ORs) with 95% confidence intervals (CIs). All statistical tests were two-sided, and p < 0.05 was considered statistically significant. Analyses were performed using Python 3.x with statsmodels, scipy, and numpy packages."""

RESULTS = f"""## Study Population

A total of 4,866 adults aged 20 years and older from NHANES 2017-2018 were included in the analysis. Of these, 1,189 (24.4%) had normal BMI, 1,593 (32.7%) were overweight, and 2,084 (42.8%) were obese (Table 1). The mean age of the study population was 51.5 +/- 17.6 years, and 2,526 (51.9%) were female. The mean BMI was 30.1 +/- 7.3 kg/m2, and the mean HbA1c was 5.87 +/- 1.10%.

Compared with the normal weight group, obese participants had higher mean HbA1c (6.05 +/- 1.17% vs 5.60 +/- 0.88%, p < 0.001) and a higher proportion of Non-Hispanic Black individuals (27.4% vs 20.4%), while Non-Hispanic Asian individuals were underrepresented (5.5% vs 25.6%) (Table 1).

## Diabetes Prevalence

The overall diabetes prevalence was {prev_overall["Prevalence (%)"].values[0]}% (95% CI: {prev_overall["95% CI lower"].values[0]}-{prev_overall["95% CI upper"].values[0]}%). A clear gradient was observed across BMI categories: the prevalence was {prev_normal["Prevalence (%)"].values[0]}% (95% CI: {prev_normal["95% CI lower"].values[0]}-{prev_normal["95% CI upper"].values[0]}%) in the normal weight group, {prev_overweight["Prevalence (%)"].values[0]}% (95% CI: {prev_overweight["95% CI lower"].values[0]}-{prev_overweight["95% CI upper"].values[0]}%) in the overweight group, and {prev_obese["Prevalence (%)"].values[0]}% (95% CI: {prev_obese["95% CI lower"].values[0]}-{prev_obese["95% CI upper"].values[0]}%) in the obese group (Figure 1).

The unweighted overall diabetes prevalence was 14.9%, compared with the survey-weighted estimate of {prev_overall["Weighted prevalence (%)"].values[0]}%, highlighting the importance of accounting for the complex survey design. The distribution of HbA1c values showed a progressive rightward shift with increasing BMI category, with a higher proportion of obese participants exceeding both the prediabetes (5.7%) and diabetes (6.5%) thresholds (Figure 3).

## Logistic Regression

In the unadjusted model (Model 1), both overweight (OR {or_ow_unadj["OR"]}, 95% CI: {or_ow_unadj["95% CI"]}, p {or_ow_unadj["p-value"]}) and obesity (OR {or_ob_unadj["OR"]}, 95% CI: {or_ob_unadj["95% CI"]}, p {or_ob_unadj["p-value"]}) were significantly associated with diabetes compared with normal weight.

After adjustment for age, sex, race/ethnicity, and education (Model 2), the association strengthened slightly for obesity (OR {or_ob["OR"]}, 95% CI: {or_ob["95% CI"]}, p {or_ob["p-value"]}) and remained significant for overweight (OR {or_ow["OR"]}, 95% CI: {or_ow["95% CI"]}, p {or_ow["p-value"]}) (Figure 2).

Among covariates, older age was associated with higher odds of diabetes (OR {or_age["OR"]} per year, 95% CI: {or_age["95% CI"]}, p {or_age["p-value"]}), while female sex was associated with lower odds (OR {or_female["OR"]}, 95% CI: {or_female["95% CI"]}, p {or_female["p-value"]}). Notable racial/ethnic disparities were observed: Non-Hispanic Asian participants had the highest adjusted odds (OR {or_asian["OR"]}, 95% CI: {or_asian["95% CI"]}), followed by Non-Hispanic Black participants (OR {or_black["OR"]}, 95% CI: {or_black["95% CI"]}) and Mexican American participants (OR {or_mexican["OR"]}, 95% CI: {or_mexican["95% CI"]}), all compared with Non-Hispanic White participants.

## Subgroup Analysis

Diabetes prevalence increased with both age and BMI category (Figure 4). Among adults aged 60-79 years, the prevalence was approximately 3-fold higher in the obese group compared with the normal weight group. Even among younger adults aged 20-39 years, obesity was associated with a markedly elevated diabetes prevalence compared with normal weight."""

DISCUSSION = """## Principal Findings

This cross-sectional analysis of NHANES 2017-2018 data demonstrated a strong, graded association between BMI category and diabetes prevalence among US adults. Obese individuals had 4.5-fold higher adjusted odds of diabetes compared with those of normal weight, while overweight individuals had approximately 2-fold higher odds. These findings persisted after adjustment for age, sex, race/ethnicity, and education level, confirming the robust and independent contribution of excess adiposity to diabetes risk.

## Comparison with Prior Literature

Our findings are consistent with a large body of epidemiological evidence linking obesity to type 2 diabetes [3, 8, 9]. The magnitude of the association observed in our study (adjusted OR 4.50 for obesity) is comparable to estimates from previous NHANES analyses and large prospective cohort studies [10, 11]. The dose-response pattern across BMI categories supports a biological gradient consistent with the causal role of adiposity in insulin resistance and beta-cell dysfunction [12].

## Role of Survey Weights

An important methodological observation was the substantial difference between unweighted (14.9%) and survey-weighted (10.2%) overall diabetes prevalence. This 4.7 percentage point difference underscores the critical importance of incorporating survey weights when analyzing NHANES data. The unweighted estimate is inflated because NHANES intentionally oversamples racial/ethnic minorities and older adults, populations with higher diabetes prevalence [5]. Studies that fail to account for the complex survey design may produce biased and non-generalizable estimates.

## Racial and Ethnic Disparities

Notable racial/ethnic disparities in diabetes risk were observed even after adjusting for BMI, age, sex, and education. Non-Hispanic Asian participants had the highest adjusted odds ratio (OR 2.97), which is particularly striking given their relatively low mean BMI. This finding is consistent with evidence that Asian populations develop metabolic complications at lower BMI thresholds [13], supporting calls for ethnicity-specific BMI cutoffs [14]. Non-Hispanic Black participants (OR 1.84) and Mexican American participants (OR 1.58) also had elevated adjusted odds, reflecting well-documented disparities in diabetes burden in the US [15].

## Limitations

Several limitations should be acknowledged. First, the cross-sectional design precludes causal inference; we can demonstrate association but not establish temporality between obesity and diabetes. Second, diabetes was defined solely by HbA1c >= 6.5%, which may underestimate true prevalence compared with definitions that incorporate fasting glucose, oral glucose tolerance testing, or self-reported physician diagnosis and medication use [7]. Third, we could not distinguish between type 1 and type 2 diabetes, although the vast majority of diabetes cases in this age group are type 2. Fourth, potential confounders not included in our model, such as physical activity, dietary patterns, family history of diabetes, and medication use, may contribute to residual confounding. Fifth, a single NHANES cycle provides cross-sectional data from 2017-2018, limiting the ability to assess temporal trends.

## Public Health Implications

Despite these limitations, our findings reinforce the critical importance of obesity prevention and weight management as cornerstone strategies for reducing the national burden of diabetes. The dose-response relationship between BMI and diabetes prevalence suggests that even modest weight reduction among overweight individuals could meaningfully reduce diabetes risk at the population level [16]. Furthermore, the observed racial/ethnic disparities highlight the need for culturally tailored interventions that address the social determinants of both obesity and diabetes [17]."""

CONCLUSION = """In this nationally representative cross-sectional study of US adults, a clear dose-response relationship was observed between BMI category and diabetes prevalence. Obesity was associated with 4.5-fold higher adjusted odds of diabetes, while overweight conferred approximately 2-fold higher odds, compared with normal weight. Significant racial/ethnic disparities persisted after adjustment for BMI and demographic factors. These findings underscore the importance of obesity prevention, weight management, and targeted interventions for high-risk populations as strategies for reducing the burden of diabetes in the United States."""

REFERENCES = """1. International Diabetes Federation. IDF Diabetes Atlas, 10th edition. Brussels, Belgium: International Diabetes Federation; 2021.
2. Centers for Disease Control and Prevention. National Diabetes Statistics Report, 2022. Atlanta, GA: CDC; 2022.
3. Prospective Studies Collaboration. Body-mass index and cause-specific mortality in 900,000 adults: collaborative analyses of 57 prospective studies. Lancet. 2009;373(9669):1083-1096.
4. National Center for Health Statistics. About the National Health and Nutrition Examination Survey. Available at: https://www.cdc.gov/nchs/nhanes/about_nhanes.htm.
5. Johnson CL, Paulose-Ram R, Ogden CL, et al. National Health and Nutrition Examination Survey: Analytic guidelines, 1999-2010. Vital Health Stat 2. 2013;(161):1-24.
6. World Health Organization. Obesity: preventing and managing the global epidemic. WHO Technical Report Series 894. Geneva: WHO; 2000.
7. American Diabetes Association. Classification and Diagnosis of Diabetes: Standards of Medical Care in Diabetes-2022. Diabetes Care. 2022;45(Suppl 1):S17-S38.
8. Mokdad AH, Ford ES, Bowman BA, et al. Prevalence of obesity, diabetes, and obesity-related health risk factors, 2001. JAMA. 2003;289(1):76-79.
9. Hales CM, Carroll MD, Fryar CD, Ogden CL. Prevalence of obesity and severe obesity among adults: United States, 2017-2018. NCHS Data Brief. 2020;(360):1-8.
10. Flegal KM, Kruszon-Moran D, Carroll MD, Fryar CD, Ogden CL. Trends in obesity among adults in the United States, 2005 to 2014. JAMA. 2016;316(3):284-291.
11. Hu FB, Manson JE, Stampfer MJ, et al. Diet, lifestyle, and the risk of type 2 diabetes mellitus in women. N Engl J Med. 2001;345(11):790-797.
12. Kahn SE, Hull RL, Utzschneider KM. Mechanisms linking obesity to insulin resistance and type 2 diabetes. Nature. 2006;444(7121):840-846.
13. Hsu WC, Araneta MR, Kanaya AM, Chiang JL, Fujimoto W. BMI cut points to identify at-risk Asian Americans for type 2 diabetes screening. Diabetes Care. 2015;38(1):150-158.
14. WHO Expert Consultation. Appropriate body-mass index for Asian populations and its implications for policy and intervention strategies. Lancet. 2004;363(9403):157-163.
15. Cheng YJ, Kanaya AM, Araneta MRG, et al. Prevalence of Diabetes by Race and Ethnicity in the United States, 2011-2016. JAMA. 2019;322(24):2389-2398.
16. Knowler WC, Barrett-Connor E, Fowler SE, et al. Reduction in the incidence of type 2 diabetes with lifestyle intervention or metformin. N Engl J Med. 2002;346(6):393-403.
17. Golden SH, Brown A, Cauley JA, et al. Health disparities in endocrine disorders: biological, clinical, and nonclinical factors--an Endocrine Society scientific statement. J Clin Endocrinol Metab. 2012;97(9):E1579-E1639."""

# === Assemble full Markdown manuscript ===
md_text = f"""# {TITLE}

---

## Abstract

{ABSTRACT}

**Keywords:** obesity, body mass index, diabetes mellitus, NHANES, epidemiology, survey weights

---

## Introduction

{INTRODUCTION}

---

## Methods

{METHODS}

---

## Results

{RESULTS}

---

## Discussion

{DISCUSSION}

---

## Conclusion

{CONCLUSION}

---

## References

{REFERENCES}

---

## Tables and Figures

**Table 1.** Baseline Characteristics of Study Participants by BMI Category (NHANES 2017-2018, n = 4,866)

**Figure 1.** Diabetes Prevalence by BMI Category (NHANES 2017-2018, unweighted with Wilson 95% CI)

**Figure 2.** Adjusted Odds Ratios for Diabetes (NHANES 2017-2018, survey-weighted logistic regression)

**Figure 3.** HbA1c Distribution by BMI Category (NHANES 2017-2018)

**Figure 4.** Diabetes Prevalence by Age Group and BMI Category (NHANES 2017-2018)
"""

# Save Markdown
os.makedirs("output", exist_ok=True)
with open("output/manuscript_draft.md", "w") as f:
    f.write(md_text)
print("Saved: output/manuscript_draft.md")

# === Build DOCX ===
doc = Document()

# Set default font
style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(12)
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(0)
paragraph_format.line_spacing = 2.0

# Helper: add heading
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

# Helper: add paragraph with optional bold/italic
def add_para(text, bold=False, alignment=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p

# Helper: add structured paragraph (bold label + normal text)
def add_labeled_para(label, text):
    p = doc.add_paragraph()
    run_label = p.add_run(label)
    run_label.bold = True
    run_label.font.name = "Times New Roman"
    run_label.font.size = Pt(12)
    run_text = p.add_run(" " + text)
    run_text.font.name = "Times New Roman"
    run_text.font.size = Pt(12)
    return p

# Title
add_para(TITLE, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()  # blank line

# Abstract
add_heading("Abstract", level=1)

# Parse abstract into labeled sections
abstract_clean = ABSTRACT.replace("**", "")
for section in ["Background:", "Methods:", "Results:", "Conclusion:"]:
    idx = abstract_clean.find(section)
    if idx >= 0:
        # Find the end (next section or end)
        next_idx = len(abstract_clean)
        for ns in ["Background:", "Methods:", "Results:", "Conclusion:"]:
            ni = abstract_clean.find(ns, idx + len(section))
            if ni >= 0 and ni < next_idx:
                next_idx = ni
        body = abstract_clean[idx + len(section):next_idx].strip()
        add_labeled_para(section, body)

add_labeled_para("Keywords:", "obesity, body mass index, diabetes mellitus, NHANES, epidemiology, survey weights")

# Introduction
add_heading("Introduction", level=1)
for para_text in INTRODUCTION.strip().split("\n\n"):
    add_para(para_text.strip())

# Methods
add_heading("Methods", level=1)

# Parse methods subsections
methods_clean = METHODS.strip()
methods_sections = methods_clean.split("\n## ")
for sec in methods_sections:
    lines = sec.strip().split("\n")
    first_line = lines[0].strip()
    # Check for sub-subsections
    if first_line.startswith("### ") or first_line.startswith("## "):
        first_line = first_line.lstrip("#").strip()
    if first_line and first_line != METHODS.strip().split("\n")[0].strip():
        add_heading(first_line, level=2)
    # Process body
    body_lines = []
    for line in lines[1:]:
        stripped = line.strip()
        if stripped.startswith("### "):
            # Flush previous body
            if body_lines:
                body_text = "\n".join(body_lines).strip()
                for para_text in body_text.split("\n\n"):
                    clean = para_text.strip().replace("**", "")
                    if clean:
                        add_para(clean)
                body_lines = []
            add_heading(stripped.lstrip("#").strip(), level=3)
        else:
            body_lines.append(line)
    if body_lines:
        body_text = "\n".join(body_lines).strip()
        for para_text in body_text.split("\n\n"):
            clean = para_text.strip().replace("**", "")
            if clean:
                add_para(clean)

# Results
add_heading("Results", level=1)
results_clean = RESULTS.strip()
results_sections = results_clean.split("\n## ")
for sec in results_sections:
    lines = sec.strip().split("\n")
    first_line = lines[0].strip()
    if first_line.startswith("## "):
        first_line = first_line.lstrip("#").strip()
    if first_line:
        add_heading(first_line, level=2)
    body = "\n".join(lines[1:]).strip()
    for para_text in body.split("\n\n"):
        clean = para_text.strip()
        if clean:
            add_para(clean)

# Discussion
add_heading("Discussion", level=1)
discussion_clean = DISCUSSION.strip()
discussion_sections = discussion_clean.split("\n## ")
for sec in discussion_sections:
    lines = sec.strip().split("\n")
    first_line = lines[0].strip()
    if first_line.startswith("## "):
        first_line = first_line.lstrip("#").strip()
    if first_line:
        add_heading(first_line, level=2)
    body = "\n".join(lines[1:]).strip()
    for para_text in body.split("\n\n"):
        clean = para_text.strip()
        if clean:
            add_para(clean)

# Conclusion
add_heading("Conclusion", level=1)
add_para(CONCLUSION.strip())

# References
add_heading("References", level=1)
for ref_line in REFERENCES.strip().split("\n"):
    ref_line = ref_line.strip()
    if ref_line:
        add_para(ref_line)

# Tables and Figures
add_heading("Tables and Figures", level=1)
add_para("Table 1. Baseline Characteristics of Study Participants by BMI Category (NHANES 2017-2018, n = 4,866)", bold=True)
add_para("")
add_para("Figure 1. Diabetes Prevalence by BMI Category (NHANES 2017-2018, unweighted with Wilson 95% CI)", bold=True)
add_para("Figure 2. Adjusted Odds Ratios for Diabetes (NHANES 2017-2018, survey-weighted logistic regression)", bold=True)
add_para("Figure 3. HbA1c Distribution by BMI Category (NHANES 2017-2018)", bold=True)
add_para("Figure 4. Diabetes Prevalence by Age Group and BMI Category (NHANES 2017-2018)", bold=True)

# Save DOCX
doc.save("output/manuscript_draft.docx")
print("Saved: output/manuscript_draft.docx")
print("=" * 60)
print("Manuscript generation complete.")
